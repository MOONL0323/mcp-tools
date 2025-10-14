"""
增强的文档解析器 - 支持多种文档格式
支持: md, txt, docx, doc, pdf, xlsx, xls, csv, 以及各种代码文件
"""

import os
import io
from pathlib import Path
from typing import Optional, Tuple
import structlog
import chardet

logger = structlog.get_logger(__name__)

# 支持的文件格式
ALLOWED_EXTENSIONS = {
    # 文档格式
    '.md', '.txt', '.doc', '.docx', '.pdf',
    # 表格格式
    '.xlsx', '.xls', '.csv',
    # 代码格式
    '.py', '.js', '.ts', '.tsx', '.jsx',
    '.java', '.go', '.cpp', '.c', '.h', '.hpp',
    '.rs', '.rb', '.php', '.swift', '.kt',
    # 配置格式
    '.json', '.yaml', '.yml', '.xml', '.toml', '.ini',
    # 其他
    '.sql', '.sh', '.bash', '.ps1', '.bat'
}

MIME_TYPE_MAPPING = {
    '.md': 'text/markdown',
    '.txt': 'text/plain',
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc': 'application/msword',
    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    '.xls': 'application/vnd.ms-excel',
    '.csv': 'text/csv',
    '.json': 'application/json',
    '.xml': 'application/xml',
}


class EnhancedDocumentParser:
    """增强的文档解析器"""
    
    @staticmethod
    def is_allowed_file(filename: str) -> bool:
        """检查文件扩展名是否支持"""
        ext = Path(filename).suffix.lower()
        return ext in ALLOWED_EXTENSIONS
    
    @staticmethod
    def get_mime_type(filename: str) -> str:
        """获取文件的MIME类型"""
        ext = Path(filename).suffix.lower()
        return MIME_TYPE_MAPPING.get(ext, 'application/octet-stream')
    
    @staticmethod
    def detect_encoding(file_content: bytes) -> str:
        """检测文件编码"""
        try:
            result = chardet.detect(file_content)
            encoding = result['encoding']
            confidence = result['confidence']
            
            logger.info(f"检测到编码: {encoding}, 置信度: {confidence}")
            
            # 如果置信度太低，使用UTF-8
            if confidence < 0.7:
                return 'utf-8'
            
            # 统一中文编码
            if encoding and encoding.lower() in ['gb2312', 'gbk', 'gb18030']:
                return 'gbk'
            
            return encoding or 'utf-8'
        except Exception as e:
            logger.warning(f"编码检测失败: {e}, 使用UTF-8")
            return 'utf-8'
    
    @staticmethod
    async def parse_file(file_path: str, original_content: bytes = None) -> Tuple[str, str]:
        """
        根据文件类型解析文档
        
        Args:
            file_path: 文件路径
            original_content: 原始文件内容（可选，用于编码检测）
            
        Returns:
            (文档文本内容, MIME类型)
        """
        file_ext = Path(file_path).suffix.lower()
        filename = Path(file_path).name
        
        # 获取MIME类型
        mime_type = EnhancedDocumentParser.get_mime_type(filename)
        
        # 根据文件类型选择解析器
        parsers = {
            '.txt': EnhancedDocumentParser._parse_text,
            '.md': EnhancedDocumentParser._parse_markdown,
            '.pdf': EnhancedDocumentParser._parse_pdf,
            '.docx': EnhancedDocumentParser._parse_docx,
            '.doc': EnhancedDocumentParser._parse_doc,
            '.xlsx': EnhancedDocumentParser._parse_excel,
            '.xls': EnhancedDocumentParser._parse_excel,
            '.csv': EnhancedDocumentParser._parse_csv,
            '.json': EnhancedDocumentParser._parse_json,
            '.yaml': EnhancedDocumentParser._parse_text,
            '.yml': EnhancedDocumentParser._parse_text,
            '.xml': EnhancedDocumentParser._parse_text,
        }
        
        # 代码文件统一用代码解析器
        code_extensions = {'.py', '.js', '.ts', '.tsx', '.jsx', '.java', '.go', 
                          '.cpp', '.c', '.h', '.hpp', '.rs', '.rb', '.php', 
                          '.swift', '.kt', '.sql', '.sh', '.bash', '.ps1', '.bat'}
        
        if file_ext in code_extensions:
            parser = EnhancedDocumentParser._parse_code
        else:
            parser = parsers.get(file_ext, EnhancedDocumentParser._parse_text)
        
        try:
            content = await parser(file_path, original_content)
            logger.info(f"成功解析文档: {filename}, 类型: {mime_type}, 长度: {len(content)}")
            return content, mime_type
        except Exception as e:
            logger.error(f"解析文档失败 {filename}: {e}")
            raise
    
    @staticmethod
    async def _parse_text(file_path: str, original_content: bytes = None) -> str:
        """解析纯文本文件（支持多种编码）"""
        try:
            # 首先尝试UTF-8
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # UTF-8失败，检测编码
            if original_content is None:
                with open(file_path, 'rb') as f:
                    original_content = f.read()
            
            encoding = EnhancedDocumentParser.detect_encoding(original_content)
            
            try:
                return original_content.decode(encoding)
            except Exception as e:
                logger.warning(f"使用{encoding}解码失败: {e}, 尝试其他编码")
                # 最后尝试使用errors='ignore'
                return original_content.decode('utf-8', errors='ignore')
    
    @staticmethod
    async def _parse_markdown(file_path: str, original_content: bytes = None) -> str:
        """解析Markdown文件"""
        content = await EnhancedDocumentParser._parse_text(file_path, original_content)
        # 保留原始markdown格式
        return content
    
    @staticmethod
    async def _parse_pdf(file_path: str, original_content: bytes = None) -> str:
        """解析PDF文件（使用pdfplumber，效果更好）"""
        try:
            import pdfplumber
            
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- 第 {page_num} 页 ---\n{text}")
            
            return '\n\n'.join(text_content)
        except ImportError:
            # 如果pdfplumber未安装，回退到PyPDF2
            logger.warning("pdfplumber未安装，使用PyPDF2解析PDF")
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- 第 {page_num} 页 ---\n{text}")
            
            return '\n\n'.join(text_content)
    
    @staticmethod
    async def _parse_docx(file_path: str, original_content: bytes = None) -> str:
        """解析Word文档(.docx)"""
        import docx
        
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # 提取表格内容
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(' | '.join(row_data))
            if table_data:
                tables_content.append('\n'.join(table_data))
        
        content = '\n\n'.join(paragraphs)
        if tables_content:
            content += '\n\n## 表格内容\n\n' + '\n\n'.join(tables_content)
        
        return content
    
    @staticmethod
    async def _parse_doc(file_path: str, original_content: bytes = None) -> str:
        """解析旧版Word文档(.doc)"""
        try:
            # 尝试使用python-docx（某些.doc文件也能读取）
            return await EnhancedDocumentParser._parse_docx(file_path, original_content)
        except Exception as e:
            logger.warning(f"无法用docx解析.doc文件: {e}")
            # 返回提示信息
            return f"""# 注意：.doc格式文件解析受限

由于技术限制，旧版.doc格式可能无法完整解析。
建议将文件转换为.docx格式后重新上传。

文件: {Path(file_path).name}
"""
    
    @staticmethod
    async def _parse_excel(file_path: str, original_content: bytes = None) -> str:
        """解析Excel文件(.xlsx/.xls)"""
        import pandas as pd
        
        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            content_parts = []
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # 转换为markdown表格格式
                content_parts.append(f"## 工作表: {sheet_name}\n")
                content_parts.append(df.to_markdown(index=False))
                content_parts.append('\n')
            
            return '\n'.join(content_parts)
        except Exception as e:
            logger.error(f"解析Excel失败: {e}")
            # 尝试简单读取
            df = pd.read_excel(file_path)
            return df.to_string()
    
    @staticmethod
    async def _parse_csv(file_path: str, original_content: bytes = None) -> str:
        """解析CSV文件"""
        import pandas as pd
        
        try:
            # 检测编码
            if original_content is None:
                with open(file_path, 'rb') as f:
                    original_content = f.read()
            
            encoding = EnhancedDocumentParser.detect_encoding(original_content)
            
            # 读取CSV
            df = pd.read_csv(file_path, encoding=encoding)
            
            # 转换为markdown表格
            return f"# CSV 数据\n\n{df.to_markdown(index=False)}"
        except Exception as e:
            logger.error(f"解析CSV失败: {e}")
            # 最后尝试纯文本读取
            return await EnhancedDocumentParser._parse_text(file_path, original_content)
    
    @staticmethod
    async def _parse_json(file_path: str, original_content: bytes = None) -> str:
        """解析JSON文件"""
        import json
        
        content = await EnhancedDocumentParser._parse_text(file_path, original_content)
        
        try:
            # 格式化JSON
            data = json.loads(content)
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
            return f"```json\n{formatted}\n```"
        except:
            # 如果不是有效JSON，返回原始内容
            return content
    
    @staticmethod
    async def _parse_code(file_path: str, original_content: bytes = None) -> str:
        """解析代码文件"""
        content = await EnhancedDocumentParser._parse_text(file_path, original_content)
        
        # 添加代码文件的元信息
        file_name = Path(file_path).name
        file_ext = Path(file_path).suffix.lstrip('.')
        
        # 使用markdown代码块格式
        return f"""# 代码文件: {file_name}

```{file_ext}
{content}
```
"""


# 向后兼容的别名
DocumentParser = EnhancedDocumentParser
